from docx import Document
import shutil

def copy_and_replace_text_in_docx(template_path, output_path, replacements):
    # Copy template to a new file
    shutil.copy(template_path, output_path)

    # Load the copied document
    doc = Document(output_path)

    # Replace text in paragraphs
    for paragraph in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in paragraph.text:
                for run in paragraph.runs:
                    run.text = run.text.replace(old_text, new_text)

    # Replace text in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_in_cell(cell, replacements)

    # Save the changes to the new document
    doc.save(output_path)

def replace_text_in_cell(cell, replacements):
    for old_text, new_text in replacements.items():
        if old_text in cell.text:
            cell.text = cell.text.replace(old_text, new_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.text = run.text.replace(old_text, new_text)

